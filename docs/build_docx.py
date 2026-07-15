"""
Convertit guide_demo.md → guide_demo.docx avec mise en forme Word complète.
Usage : python docs/build_docx.py
"""
from __future__ import annotations

import os
import re
import sys

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, BASE_DIR)

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_PATH  = os.path.join(os.path.dirname(__file__), "guide_demo.md")
OUT_PATH = os.path.join(os.path.dirname(__file__), "guide_demo.docx")

# ── Palette couleurs ─────────────────────────────────────────────────
C_TITLE   = RGBColor(0x1F, 0x35, 0x64)   # bleu foncé
C_H2      = RGBColor(0x1F, 0x35, 0x64)
C_H3      = RGBColor(0x2E, 0x74, 0xB5)   # bleu moyen
C_H4      = RGBColor(0x2E, 0x74, 0xB5)
C_CODE_BG = RGBColor(0xF2, 0xF2, 0xF2)
C_CODE_FG = RGBColor(0x2E, 0x2E, 0x2E)
C_OK      = RGBColor(0x37, 0x86, 0x37)   # vert
C_WARN    = RGBColor(0xC5, 0x5A, 0x11)   # orange

# ── Helpers ──────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Colorie le fond d'une cellule de tableau."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = C_TITLE
        run.font.size = Pt(22)
    elif level == 2:
        run.font.color.rgb = C_H2
        run.font.size = Pt(16)
    elif level == 3:
        run.font.color.rgb = C_H3
        run.font.size = Pt(13)
    else:
        run.font.color.rgb = C_H4
        run.font.size = Pt(12)
    return p

def add_code_block(doc: Document, lines: list[str]):
    """Bloc de code avec fond gris et police monospace."""
    text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # fond gris via shading
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = C_CODE_FG

def add_table_from_md(doc: Document, rows_md: list[str]):
    """Parse un tableau Markdown et l'insère en Word."""
    data_rows = [r for r in rows_md if not re.match(r"^\s*\|[-| :]+\|\s*$", r)]
    if not data_rows:
        return
    cells = [[c.strip() for c in row.strip().strip("|").split("|")] for row in data_rows]
    ncols = max(len(r) for r in cells)
    table = doc.add_table(rows=len(cells), cols=ncols)
    table.style = "Table Grid"
    for i, row_data in enumerate(cells):
        for j, val in enumerate(row_data):
            if j >= ncols:
                break
            cell = table.cell(i, j)
            cell.text = val
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if i == 0:
                run.bold = True
                set_cell_bg(cell, "1F3564")
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            else:
                bg = "DDEEFF" if i % 2 == 0 else "FFFFFF"
                set_cell_bg(cell, bg)
    doc.add_paragraph()

def inline_format(para, text: str):
    """Applique le gras/italique inline dans un paragraphe."""
    parts = re.split(r"(\*\*.*?\*\*|`[^`]+`)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = C_CODE_FG
        else:
            para.add_run(part)


# ── Parser principal ─────────────────────────────────────────────────

def build_docx(md_path: str, out_path: str):
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Style Normal
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    lines = open(md_path, encoding="utf-8").read().splitlines()

    in_code  = False
    code_buf: list[str] = []
    table_buf: list[str] = []
    in_table = False

    i = 0
    while i < len(lines):
        line = lines[i]

        # ── Bloc de code ────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                add_code_block(doc, code_buf)
                in_code  = False
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── Tableau Markdown ────────────────────────────────────────
        if line.strip().startswith("|"):
            table_buf.append(line)
            i += 1
            continue
        elif table_buf:
            add_table_from_md(doc, table_buf)
            table_buf = []

        # ── Titres ──────────────────────────────────────────────────
        if line.startswith("#### "):
            add_heading(doc, line[5:], 4)
        elif line.startswith("### "):
            add_heading(doc, line[4:], 3)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 2)
        elif line.startswith("# "):
            add_heading(doc, line[2:], 1)

        # ── Ligne de séparation ─────────────────────────────────────
        elif re.match(r"^-{3,}$", line.strip()):
            doc.add_paragraph("─" * 60)

        # ── Liste à puces ───────────────────────────────────────────
        elif line.strip().startswith("- ") or line.strip().startswith("* "):
            text = re.sub(r"^[\s\-\*]+", "", line)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.5)
            inline_format(p, text)

        # ── Liste numérotée ─────────────────────────────────────────
        elif re.match(r"^\d+\. ", line.strip()):
            text = re.sub(r"^\d+\.\s+", "", line.strip())
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Cm(0.5)
            inline_format(p, text)

        # ── Ligne vide ──────────────────────────────────────────────
        elif line.strip() == "":
            doc.add_paragraph()

        # ── Paragraphe normal ────────────────────────────────────────
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            inline_format(p, line)

        i += 1

    # Vider les buffers résiduels
    if table_buf:
        add_table_from_md(doc, table_buf)
    if code_buf:
        add_code_block(doc, code_buf)

    doc.save(out_path)
    print(f"  ✔  DOCX généré : {out_path}  ({os.path.getsize(out_path) // 1024} Ko)")


if __name__ == "__main__":
    build_docx(MD_PATH, OUT_PATH)
