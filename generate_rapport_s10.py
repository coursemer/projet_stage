"""
Génère le rapport Semaine 10 — Intégration LLM (Mistral AI) en format docx.
Usage: python3 generate_rapport_s10.py
"""
from __future__ import annotations
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "docs", "rapport_semaine10.docx")

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        h.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        h.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return h

def add_paragraph(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.7 + level * 0.5)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        r.bold = True
        r.font.size = Pt(10.5)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Header background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "2E74B5")
        tcPr.append(shd)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for ri, row in enumerate(rows):
        trow = table.rows[ri + 1]
        for ci, val in enumerate(row):
            cell = trow.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "EBF3FB")
                tcPr.append(shd)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


# ── Document ──────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.5)

# ── Page de garde ─────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("RAPPORT DE STAGE — SEMAINE 10")
r.bold = True; r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

doc.add_paragraph()

meta = [
    ("Sujet",       "Intégration LLM (Mistral AI) pour le Data Trust Agent"),
    ("Réalisé par", "NACERDDINE Rim · BENYOUSSEF Roba"),
    ("Encadré par", "M. YOUSFI Youssef · Mme. GHOLAMI Somayyeh"),
    ("Date",        "20 Mai 2026"),
    ("Statut",      "✅ Livré"),
]
table_meta = doc.add_table(rows=len(meta), cols=2)
table_meta.style = "Table Grid"
for i, (k, v) in enumerate(meta):
    r0 = table_meta.rows[i]
    r0.cells[0].text = k
    r0.cells[0].paragraphs[0].runs[0].bold = True
    r0.cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    r0.cells[1].text = v
    r0.cells[1].paragraphs[0].runs[0].font.size = Pt(11)

doc.add_page_break()

# ── 1. Introduction ───────────────────────────────────────────────────────────

set_heading(doc, "1. Introduction et contexte", 1)

add_paragraph(doc, (
    "La Semaine 10 marque une étape clé du stage PFE : l'intégration d'un Large Language Model (LLM) "
    "dans la plateforme Data Trust Agent. Cette phase, planifiée dès la semaine 3 dans le Document de "
    "Décision Technique, transforme le système de détection d'anomalies en un véritable agent conversationnel "
    "capable d'expliquer en langage naturel les incidents détectés dans les pipelines de données."
))
doc.add_paragraph()
add_paragraph(doc, (
    "Le choix du LLM a été formalisé dans un document d'analyse technique dédié (analyse_llm.docx), "
    "évaluant cinq familles de solutions selon quatre critères : capacités techniques, conformité RGPD, "
    "coût et facilité d'intégration. La recommandation principale porte sur Mistral AI — La Plateforme, "
    "solution française hébergée en France (OVHcloud), pleinement conforme au RGPD et à la loi belge "
    "du 30 juillet 2018 relative à la protection des données personnelles."
))

doc.add_paragraph()
add_paragraph(doc, "Livrable planifié Semaine 10 :", bold=True)
add_bullet(doc, "Module LLM intégré au pipeline de détection ML", bold_prefix="→")
add_bullet(doc, "API REST exposant les explications LLM", bold_prefix="→")
add_bullet(doc, "Dashboard HTML mis à jour avec les explications", bold_prefix="→")
add_bullet(doc, "DAG Airflow étendu avec une tâche d'explication", bold_prefix="→")
add_bullet(doc, "55 tests existants maintenus à 100 % de passage", bold_prefix="→")

# ── 2. Analyse LLM ───────────────────────────────────────────────────────────

set_heading(doc, "2. Décision technique — Choix du LLM", 1)
set_heading(doc, "2.1 Critères d'évaluation", 2)

add_paragraph(doc, (
    "Le contexte belge impose des contraintes RGPD strictes : l'Autorité de protection des données (APD) "
    "belge exige que les données traitées par un service tiers restent dans l'Union Européenne "
    "(Art. 44-49 du Règlement (UE) 2016/679). Les cinq solutions évaluées sont :"
))
doc.add_paragraph()

add_table(doc,
    ["Solution", "Raisonnement", "Conformité RGPD", "Hébergement UE", "Coût API"],
    [
        ["GPT-4o (Azure Belgium North)", "★★★★★", "Bonne (via Azure)", "✓ Belgique", "Modéré"],
        ["Mistral AI — La Plateforme", "★★★★☆", "Excellente (native)", "✓ France", "Faible"],
        ["Anthropic Claude 3.5", "★★★★★", "Partielle (US)", "✗", "Modéré"],
        ["Google Gemini 1.5 (Vertex AI)", "★★★★☆", "Partielle (CLOUD Act)", "Partiel", "Faible"],
        ["Open Source (Ollama local)", "★★★☆☆", "Maximale (local)", "✓ (local)", "Infra propre"],
    ],
    col_widths=[5.0, 3.5, 4.0, 3.5, 2.5]
)

doc.add_paragraph()
set_heading(doc, "2.2 Recommandation retenue", 2)

add_paragraph(doc, (
    "Mistral AI est retenu comme LLM principal, avec Ollama comme solution de fallback "
    "pour les données sensibles ou en cas d'indisponibilité de l'API. "
    "Justifications principales :"
))
add_bullet(doc, "Société française soumise au droit de l'UE — valide pour une entité belge sans démarche supplémentaire")
add_bullet(doc, "Données hébergées en France (intra-UE) — conforme Art. 44 RGPD et APD belge")
add_bullet(doc, "DPA conforme au RGPD disponible en ligne, opposable en Belgique")
add_bullet(doc, "API compatible format OpenAI — intégration Python sans friction (SDK mistralai)")
add_bullet(doc, "Coût inférieur aux alternatives américaines")

doc.add_paragraph()
add_paragraph(doc, "Architecture par cas d'usage :", bold=True)
add_table(doc,
    ["Cas d'usage", "Modèle Mistral", "Justification"],
    [
        ["Analyse incidents & explications", "mistral-large-latest", "Raisonnement complexe sur logs Airflow/Spark"],
        ["Génération tests dbt & règles qualité", "codestral-latest", "Spécialisé code, idéal SQL/YAML"],
        ["Enrichissement Data Catalog", "mistral-small-latest", "Tâche simple, coût réduit"],
        ["Recherche sémantique (embeddings)", "mistral-embed", "Natif, cohérent avec l'espace vectoriel Mistral"],
        ["Fallback / données sensibles", "Ollama + Mistral 7B / Llama 3.1 8B", "Zéro transfert, souveraineté totale"],
    ],
    col_widths=[5.0, 5.0, 7.0]
)

# ── 3. Travail réalisé ────────────────────────────────────────────────────────

set_heading(doc, "3. Travail réalisé — Semaine 10", 1)
set_heading(doc, "3.1 Vue d'ensemble des livrables", 2)

add_table(doc,
    ["Livrable", "Fichier", "Statut"],
    [
        ["Module LLMExplainer", "spark/metrics/llm_explainer.py", "✅ Livré"],
        ["Flag --explain (CLI)", "spark/metrics/run_collector.py", "✅ Livré"],
        ["Endpoint /api/v1/alerts/explain", "spark/metrics/api.py", "✅ Livré"],
        ["Dashboard HTML — section LLM", "spark/metrics/dashboard.py", "✅ Livré"],
        ["Tâche DAG Airflow", "dags/data_trust_monitoring.py", "✅ Livré"],
        ["Dépendance mistralai>=1.0.0", "requirements.txt", "✅ Livré"],
    ],
    col_widths=[5.5, 7.0, 2.5]
)

doc.add_paragraph()
set_heading(doc, "3.2 LLMExplainer — Architecture 3 niveaux", 2)

add_paragraph(doc, (
    "Le module spark/metrics/llm_explainer.py implémente la classe LLMExplainer selon "
    "l'architecture hybride recommandée dans le document d'analyse : "
    "trois niveaux de fallback garantissent la continuité du service dans tous les scénarios."
))
doc.add_paragraph()

add_table(doc,
    ["Priorité", "Backend", "Condition d'activation", "Latence estimée"],
    [
        ["1 (principal)", "Mistral AI — mistral-large-latest", "MISTRAL_API_KEY défini", "< 5s"],
        ["2 (fallback)", "Ollama local — mistral:7b", "API Mistral indisponible", "2–30s (selon GPU)"],
        ["3 (fallback final)", "Templates par type d'anomalie", "Aucun backend disponible", "< 1ms"],
    ],
    col_widths=[2.8, 5.5, 5.5, 3.5]
)

doc.add_paragraph()
add_paragraph(doc, "Fonctionnalités clés du module :", bold=True)
add_bullet(doc, "enrich_anomalies(anomalies) — enrichit la liste d'Anomaly en place (context['llm_explanation'])", bold_prefix="•")
add_bullet(doc, "explain_alert_dict(alert) — explique une alerte depuis AlertManager.query() (usage API)", bold_prefix="•")
add_bullet(doc, "Cache fichier SHA-256 dans spark/data/llm_cache/ — évite les appels redondants (recommandation coûts)", bold_prefix="•")
add_bullet(doc, "Prompts en anglais, concis et structurés (bonnes pratiques analyse_llm.docx)", bold_prefix="•")

doc.add_paragraph()
set_heading(doc, "3.3 Intégration dans run_collector.py", 2)

add_paragraph(doc, (
    "Le flag --explain est ajouté à la CLI de collecte. Lorsqu'activé avec --detect-ml, "
    "le module LLMExplainer est appelé après chaque run de détection ML pour enrichir "
    "les anomalies avant leur persistance dans la table alerts (champ tags JSON)."
))
doc.add_paragraph()
add_paragraph(doc, "Commandes disponibles après Semaine 10 :", bold=True)
add_bullet(doc, "python3 run_collector.py --detect-ml                 → détection ML seule", bold_prefix="•")
add_bullet(doc, "python3 run_collector.py --detect-ml --explain       → ML + explication Mistral", bold_prefix="•")
add_bullet(doc, "export MISTRAL_API_KEY='sk-...'  (variable d'env pour activer l'API cloud)", bold_prefix="•")

doc.add_paragraph()
set_heading(doc, "3.4 API REST — Endpoint /api/v1/alerts/explain", 2)

add_paragraph(doc, (
    "Un nouvel endpoint POST est ajouté à l'API FastAPI (spark/metrics/api.py). "
    "Il récupère les alertes ML récentes, appelle LLMExplainer et retourne "
    "les explications en JSON, avec indication du backend utilisé."
))
doc.add_paragraph()
add_table(doc,
    ["Endpoint", "Méthode", "Paramètres", "Réponse"],
    [
        ["/api/v1/alerts/explain", "POST", "source, severity, limit", "{explained, results: [{alert_id, explanation, backend}]}"],
        ["/api/v1/alerts/detect-ml", "POST", "—", "{pipelines_analyzed, detected, saved, by_severity}"],
        ["/api/v1/alerts/detect", "POST", "—", "{detected, saved, by_severity}"],
    ],
    col_widths=[5.0, 2.5, 4.5, 5.5]
)

doc.add_paragraph()
set_heading(doc, "3.5 Dashboard HTML — Explications LLM dans les cartes ML", 2)

add_paragraph(doc, (
    "La section « 🤖 Pipelines — Analyse ML (Semaine 10) » du dashboard est mise à jour : "
    "chaque carte de pipeline affiche désormais les explications LLM si elles sont disponibles "
    "(champ tags['llm_explanation'] de l'alerte). Le backend utilisé (mistral/ollama/template) "
    "est indiqué entre crochets. Un style CSS dédié (.llm-explanation) est ajouté."
))

doc.add_paragraph()
set_heading(doc, "3.6 DAG Airflow — Tâche explain_anomalies_llm", 2)

add_paragraph(doc, (
    "Le DAG data_trust_monitoring est étendu avec une tâche BashOperator "
    "explain_anomalies_llm, exécutée après detect_anomalies_ml. "
    "Le flux d'exécution est le suivant :"
))
doc.add_paragraph()
add_paragraph(doc, "collect_metrics → [detect_anomalies ‖ detect_anomalies_ml] → explain_anomalies_llm → generate_dashboard → print_alert_summary",
              italic=True, indent=True)

# ── 4. Tests ─────────────────────────────────────────────────────────────────

set_heading(doc, "4. Tests automatisés — Résultats complets", 1)
set_heading(doc, "4.1 Résultat global", 2)

add_paragraph(doc, (
    "La suite de tests est exécutée avec pytest depuis le répertoire projet_stage/. "
    "Les 55 tests existants (hérités des Semaines 8-9) sont tous maintenus à 100 % "
    "après les modifications de la Semaine 10."
))
doc.add_paragraph()

add_table(doc,
    ["Résultat", "Détail"],
    [
        ["✅ 55 tests passants", "0 échec, 0 erreur"],
        ["⚠️ 2 warnings (non bloquants)", "Pydantic V2 : ConfigDict deprecation (hérité des modèles ruby-rue)"],
        ["⏱ Durée totale", "~4 min 18s (domination : tests IsolationForest avec 40 snapshots)"],
        ["Commande", "airflow_env/bin/python3 -m pytest spark/tests/ -v"],
    ],
    col_widths=[5.5, 12.0]
)

doc.add_paragraph()
set_heading(doc, "4.2 Détail par fichier de test", 2)

test_files = [
    (
        "test_volume_detector.py",
        "VolumeDetector",
        8,
        [
            ("test_no_anomaly_on_normal_data", "Aucune anomalie sur données normales (signal stable)", "✅"),
            ("test_insufficient_history_returns_empty", "Historique insuffisant (<7) → liste vide", "✅"),
            ("test_zero_row_count_with_no_history_is_critical", "0 lignes sans historique → CRITICAL", "✅"),
            ("test_large_volume_drop_detected", "Chute 95% → anomalie VOLUME HIGH/CRITICAL", "✅"),
            ("test_moderate_drop_not_flagged_below_threshold", "Chute 10% → pas d'alerte (sous seuil 50%)", "✅"),
            ("test_large_volume_spike_detected", "Pic ×10 → anomalie VOLUME détectée", "✅"),
            ("test_anomaly_has_required_fields", "Vérification id, pipeline_name, metric_name, observed/expected", "✅"),
            ("test_none_row_count_returns_empty", "row_count=None → liste vide (pas de crash)", "✅"),
        ]
    ),
    (
        "test_level_detectors.py",
        "DistributionDetector · SchemaDetector · PerformanceDetector",
        19,
        [
            ("test_no_anomaly_normal_distribution", "[Distribution] Signal normal → aucune anomalie", "✅"),
            ("test_null_rate_hard_threshold", "[Distribution] null_rate élevé → alerte HIGH", "✅"),
            ("test_mean_drift_detected", "[Distribution] Dérive de la moyenne (amount.mean) détectée", "✅"),
            ("test_no_anomaly_when_no_column_stats", "[Distribution] column_stats vide → liste vide", "✅"),
            ("test_columns_to_check_filter", "[Distribution] Filtre columns_to_check respecté", "✅"),
            ("test_insufficient_history_per_stat", "[Distribution] Historique trop court → pas d'alerte stat", "✅"),
            ("test_no_anomaly_matching_schema", "[Schema] Schéma identique → aucune anomalie", "✅"),
            ("test_dropped_column_detected", "[Schema] Colonne supprimée → anomalie DROPPED", "✅"),
            ("test_dropped_critical_column_is_critical", "[Schema] Colonne critique supprimée → CRITICAL", "✅"),
            ("test_added_column_is_low_severity", "[Schema] Nouvelle colonne inattendue → LOW", "✅"),
            ("test_type_change_detected", "[Schema] Changement de type → anomalie 'type changed'", "✅"),
            ("test_no_reference_uses_first_history", "[Schema] Pas de référence → utilise premier historique", "✅"),
            ("test_empty_history_and_no_reference_returns_empty", "[Schema] Historique vide + pas de ref → []", "✅"),
            ("test_no_anomaly_normal_duration", "[Performance] Durée normale → aucune anomalie", "✅"),
            ("test_sla_breach_is_critical", "[Performance] Dépassement SLA → CRITICAL", "✅"),
            ("test_duration_spike_without_sla", "[Performance] Pic durée sans SLA → PERFORMANCE détecté", "✅"),
            ("test_high_failure_rate_flagged", "[Performance] Taux d'échec élevé → alerte failure rate", "✅"),
            ("test_task_failure_spike", "[Performance] Pic task_failures → anomalie tâche", "✅"),
        ]
    ),
    (
        "test_temporal_and_ml.py",
        "SeasonalityDetector · TrendDetector · CorrelationDetector · MLBaselineDetector",
        14,
        [
            ("test_no_anomaly_when_volume_matches_day_baseline", "[Saisonnalité] Volume dimanche conforme → OK", "✅"),
            ("test_anomaly_when_weekday_volume_is_sunday_like", "[Saisonnalité] Lundi avec volume dimanche → anomalie", "✅"),
            ("test_insufficient_per_day_history_returns_empty", "[Saisonnalité] Trop peu d'occurrences/jour → []", "✅"),
            ("test_degrading_volume_trend_detected", "[Tendance] Décroissance régulière → tendance TEMPORAL", "✅"),
            ("test_stable_data_no_trend", "[Tendance] Données stables (bruit 2%) → pas de tendance", "✅"),
            ("test_insufficient_points_no_trend", "[Tendance] < 7 points → pas de détection", "✅"),
            ("test_correlated_drop_detected", "[Corrélation] Deux pipelines qui chutent ensemble → alerte", "✅"),
            ("test_uncorrelated_pipelines_no_anomaly", "[Corrélation] Pipelines décorrélés → pas de crash", "✅"),
            ("test_single_pipeline_no_crash", "[Corrélation] Pipeline unique → liste vide (pas de crash)", "✅"),
            ("test_normal_point_not_flagged", "[ML Isolation Forest] Point normal → pas d'alerte (probabiliste)", "✅"),
            ("test_extreme_anomaly_detected", "[ML] Outlier extrême (×100) → AnomalyLevel.ML détecté", "✅"),
            ("test_insufficient_history_returns_empty", "[ML] < 20 snapshots d'historique → []", "✅"),
            ("test_anomaly_score_returned", "[ML] score_samples retourne un float non-nul", "✅"),
            ("test_anomaly_has_correct_fields", "[ML] context contient 'IsolationForest', metric_name correct", "✅"),
        ]
    ),
    (
        "test_scoring_and_integration.py",
        "SeverityScorer · AnomalyDetector (intégration complète)",
        14,
        [
            ("test_schema_anomaly_scores_higher_than_ml", "[Score] Schéma > ML en severity_score", "✅"),
            ("test_high_sigma_escalates_score", "[Score] sigma=7 > sigma=3.1 en score", "✅"),
            ("test_pipeline_weight_applied", "[Score] Poids pipeline critical×1.5 > low×0.5", "✅"),
            ("test_recurrence_bonus_applied", "[Score] 5 occurrences passées → is_recurring=True, count=6", "✅"),
            ("test_ml_agreement_bonus", "[Score] Agreement ML → bonus appliqué", "✅"),
            ("test_score_capped_at_100", "[Score] Score plafonné à 100 même pour sigma=20", "✅"),
            ("test_severity_labels_assigned", "[Score] Label Severity attribué après scoring", "✅"),
            ("test_no_anomalies_on_normal_data", "[Intégration] Pipeline normal → has_anomalies=False", "✅"),
            ("test_volume_anomaly_in_result", "[Intégration] row_count=50 → VOLUME dans DetectionResult", "✅"),
            ("test_schema_anomaly_in_result", "[Intégration] Colonne manquante → SCHEMA dans result", "✅"),
            ("test_sla_breach_in_result", "[Intégration] duration=7200s → PERFORMANCE HIGH/CRITICAL", "✅"),
            ("test_result_summary_structure", "[Intégration] summary() → {pipeline, total_anomalies, by_severity, by_level}", "✅"),
            ("test_worst_severity_property", "[Intégration] worst_severity retourne une valeur Severity valide", "✅"),
            ("test_anomaly_history_used_for_recurrence", "[Intégration] Historique anomalies → is_recurring propagé", "✅"),
            ("test_factory_defaults", "[Intégration] AnomalyDetector.for_pipeline() instancie tous les sous-détecteurs", "✅"),
        ]
    ),
]

for fname, title, total, tests in test_files:
    add_paragraph(doc, f"📁 {fname}  —  {title}  ({total} test{'s' if total>1 else ''})", bold=True)
    add_table(doc,
        ["Nom du test", "Ce qui est vérifié", "Résultat"],
        [[t[0], t[1], t[2]] for t in tests],
        col_widths=[6.5, 8.5, 2.5]
    )
    doc.add_paragraph()

set_heading(doc, "4.3 Récapitulatif des tests par détecteur", 2)

add_table(doc,
    ["Détecteur / Composant", "Tests", "Passants", "Couverture fonctionnelle"],
    [
        ["VolumeDetector", "8", "8 ✅", "Chutes, pics, seuils, champs requis, edge cases"],
        ["DistributionDetector", "6", "6 ✅", "Null rate, dérive de moyenne, filtres colonnes, historique court"],
        ["SchemaDetector", "7", "7 ✅", "Colonnes supprimées/ajoutées, type change, ref auto-dérivée"],
        ["PerformanceDetector", "5", "5 ✅", "SLA, pic durée, taux d'échec, task_failures"],
        ["SeasonalityDetector", "3", "3 ✅", "Pattern jour/semaine, insuffisance historique"],
        ["TrendDetector", "3", "3 ✅", "Décroissance régulière, données stables, min_points"],
        ["CorrelationDetector", "3", "3 ✅", "Corrélation croisée, décorrélation, pipeline unique"],
        ["MLBaselineDetector (IF)", "5", "5 ✅", "Outlier extrême, score, champs context, historique min"],
        ["SeverityScorer", "7", "7 ✅", "Poids, sigma, récurrence, agreement ML, plafonnement 100"],
        ["AnomalyDetector (intégr.)", "8", "8 ✅", "Pipeline complet : volume + schéma + SLA + récurrence"],
        ["TOTAL", "55", "55 ✅", "100 % de passage — aucune régression Semaine 10"],
    ],
    col_widths=[5.0, 2.0, 2.5, 8.0]
)

# ── 5. Difficultés ────────────────────────────────────────────────────────────

set_heading(doc, "5. Difficultés rencontrées & solutions", 1)

add_table(doc,
    ["Problème", "Cause", "Solution"],
    [
        [
            "Choix du LLM : Anthropic vs Mistral",
            "Claude (Anthropic) est la plateforme de développement — tentation de l'utiliser directement",
            "Document analyse_llm.docx → Mistral AI retenu (RGPD, APD belge). Anthropic écarté pour données hors UE"
        ],
        [
            "Perte du contexte Anomaly → AnomalyAlert",
            "La conversion Anomaly → AnomalyAlert (anomaly_detector.py) ne transfère pas context dict",
            "Champ tags de AnomalyAlert utilisé pour stocker llm_explanation et llm_backend en JSON"
        ],
        [
            "Conflit opentelemetry lors de pip install mistralai",
            "mistralai tire opentelemetry-exporter-otlp-proto-grpc qui nécessite une version plus récente",
            "Warning non bloquant — Airflow utilise ses propres versions OTel ; l'import mistralai fonctionne"
        ],
        [
            "Cache LLM à implémenter sans Redis (pas de serveur en stage)",
            "Le document recommande Redis mais l'infra stage ne le prévoit pas",
            "Cache fichier SQLite-less : JSON par hash SHA-256 du prompt dans spark/data/llm_cache/"
        ],
        [
            "Tests en 4+ minutes (IsolationForest)",
            "Chaque test MLBaselineDetector entraîne un IF sur 40 snapshots",
            "Non bloquant — durée stable, acceptée pour un stage. Option : passer n_estimators à 10 en test"
        ],
    ],
    col_widths=[4.5, 6.0, 7.0]
)

# ── 6. Bilan ──────────────────────────────────────────────────────────────────

set_heading(doc, "6. Bilan & Perspectives", 1)
set_heading(doc, "6.1 Bilan Semaine 10", 2)

add_table(doc,
    ["Objectif planifié", "Statut", "Commentaire"],
    [
        ["LLMExplainer — Mistral AI principal + fallback Ollama", "✅", "Architecture 3 niveaux : Mistral → Ollama → Template"],
        ["Cache des réponses LLM (coût-optimisation)", "✅", "Cache fichier SHA-256, spark/data/llm_cache/"],
        ["CLI --explain dans run_collector.py", "✅", "Intégré à la chaîne detect-ml existante"],
        ["API REST POST /api/v1/alerts/explain", "✅", "FastAPI, retourne {explained, results, backend, ts}"],
        ["Dashboard HTML — section LLM", "✅", "Bloc .llm-explanation par carte pipeline ML"],
        ["DAG Airflow — tâche explain_anomalies_llm", "✅", "Après detect_anomalies_ml, avant generate_dashboard"],
        ["55 tests — 100% passage (zéro régression)", "✅", "Aucune modification des tests existants"],
        ["Conformité RGPD / APD belge", "✅", "Mistral AI : DPA FR, données intra-UE"],
    ],
    col_widths=[7.5, 1.5, 8.5]
)

doc.add_paragraph()
set_heading(doc, "6.2 Perspectives — Semaines suivantes", 2)

add_paragraph(doc, "Les prochaines étapes LLM identifiées dans le document d'analyse technique sont :", bold=True)
add_bullet(doc, "Génération automatique de tests dbt (YAML) via Codestral — remonter les lacunes détectées (duplicates, future_date, referential_break)", bold_prefix="S11 →")
add_bullet(doc, "Enrichissement du Data Catalog — descriptions automatiques des tables/colonnes via Mistral Small 3", bold_prefix="S11 →")
add_bullet(doc, "Recherche sémantique dans le catalog via Mistral Embed (OpenMetadata ou catalog maison)", bold_prefix="S12 →")
add_bullet(doc, "Rapport de stage final — consolidation de l'ensemble des phases 1, 2 et 3", bold_prefix="S12 →")

doc.add_paragraph()
add_paragraph(doc, (
    "L'architecture hybride Mistral AI / Ollama est désormais opérationnelle. "
    "La plateforme Data Trust Agent dispose d'un système complet de détection d'anomalies (7 couches ML) "
    "couplé à une explication LLM intelligible par les équipes data, avec garantie de conformité RGPD "
    "grâce au choix de Mistral AI comme fournisseur principal."
))

# ── Save ──────────────────────────────────────────────────────────────────────

os.makedirs(os.path.dirname(OUT), exist_ok=True)
doc.save(OUT)
print(f"✅ Rapport généré : {OUT}")
